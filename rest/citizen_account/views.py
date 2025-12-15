from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from .models import CitizenAccountEntry
from .forms import CitizenAccountEntryForm
from django.db.models import Q
import openpyxl
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
import datetime

def citizen_account_list(request):
    query = request.GET.get("q")
    entries = CitizenAccountEntry.objects.all()
    if query:
        entries = entries.filter(Q(name__icontains=query) | Q(identity__icontains=query))
    return render(request, 'citizen_account/citizen_account_list.html', {'entries': entries, 'query': query})

def citizen_account_create(request):
    if request.method == 'POST':
        form = CitizenAccountEntryForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('citizen_account_list')
    else:
        form = CitizenAccountEntryForm()
    return render(request, 'citizen_account/citizen_account_form.html', {'form': form})

def citizen_account_update(request, pk):
    entry = get_object_or_404(CitizenAccountEntry, pk=pk)
    if request.method == 'POST':
        form = CitizenAccountEntryForm(request.POST, instance=entry)
        if form.is_valid():
            form.save()
            return redirect('citizen_account_list')
    else:
        form = CitizenAccountEntryForm(instance=entry)
    return render(request, 'citizen_account/citizen_account_form.html', {'form': form})

def citizen_account_delete(request, pk):
    entry = get_object_or_404(CitizenAccountEntry, pk=pk)
    if request.method == 'POST':
        entry.delete()
        return redirect('citizen_account_list')
    return render(request, 'citizen_account/citizen_account_confirm_delete.html', {'entry': entry})

def export_citizen_account_to_excel(request):
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="citizen_account_data.xlsx"'

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Citizen Account Data"

    # Write headers
    columns = ['Name', 'Identity', 'Registration Date']
    sheet.append(columns)

    # Write data
    entries = CitizenAccountEntry.objects.all()
    for entry in entries:
        sheet.append([entry.name, entry.identity, entry.registration_date])

    workbook.save(response)
    return response

def generate_citizen_account_report(request):
    document = Document()
    document.add_heading('تقرير حساب المواطن اليومي', 0)

    today = datetime.date.today()
    document.add_paragraph(f'التاريخ: {today.strftime("%Y-%m-%d")}')

    entries = CitizenAccountEntry.objects.filter(registration_date=today)
    if entries:
        document.add_heading('مدخلات اليوم', level=1)
        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'الاسم'
        hdr_cells[1].text = 'الهوية'
        hdr_cells[2].text = 'تاريخ التسجيل'
        for entry in entries:
            row_cells = table.add_row().cells
            row_cells[0].text = entry.name
            row_cells[1].text = entry.identity
            row_cells[2].text = str(entry.registration_date)
    else:
        document.add_paragraph('لا توجد مدخلات جديدة لحساب المواطن اليوم.')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="citizen_account_report_{today.strftime("%Y-%m-%d")}.docx"'
    document.save(response)
    return response

