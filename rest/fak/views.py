from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from .models import FakEntry
from .forms import FakEntryForm
from django.db.models import Q
import openpyxl
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
import datetime

def fak_list(request):
    query = request.GET.get("q")
    entries = FakEntry.objects.all()
    if query:
        entries = entries.filter(Q(name__icontains=query) | Q(identity__icontains=query) | Q(payment_status__icontains=query) | Q(whatsapp_name__icontains=query))
    return render(request, 'fak/fak_list.html', {'entries': entries, 'query': query})

def fak_create(request):
    if request.method == 'POST':
        form = FakEntryForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('fak_list')
    else:
        form = FakEntryForm()
    return render(request, 'fak/fak_form.html', {'form': form})

def fak_update(request, pk):
    entry = get_object_or_404(FakEntry, pk=pk)
    if request.method == 'POST':
        form = FakEntryForm(request.POST, instance=entry)
        if form.is_valid():
            form.save()
            return redirect('fak_list')
    else:
        form = FakEntryForm(instance=entry)
    return render(request, 'fak/fak_form.html', {'form': form})

def fak_delete(request, pk):
    entry = get_object_or_404(FakEntry, pk=pk)
    if request.method == 'POST':
        entry.delete()
        return redirect('fak_list')
    return render(request, 'fak/fak_confirm_delete.html', {'entry': entry})

def export_fak_to_excel(request):
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="fak_data.xlsx"'

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Fak Data"

    # Write headers
    columns = ['Name', 'Identity', 'Payment Status', 'WhatsApp Name', 'Entry Date']
    sheet.append(columns)

    # Write data
    entries = FakEntry.objects.all()
    for entry in entries:
        sheet.append([entry.name, entry.identity, entry.payment_status, entry.whatsapp_name, entry.entry_date])

    workbook.save(response)
    return response

def generate_fak_report(request):
    document = Document()
    document.add_heading('تقرير فك اليومي', 0)

    today = datetime.date.today()
    document.add_paragraph(f'التاريخ: {today.strftime("%Y-%m-%d")}')

    entries = FakEntry.objects.filter(entry_date=today)
    if entries:
        document.add_heading('مدخلات اليوم', level=1)
        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'الاسم'
        hdr_cells[1].text = 'الهوية'
        hdr_cells[2].text = 'حالة السداد'
        hdr_cells[3].text = 'الاسم في الواتساب'
        hdr_cells[4].text = 'تاريخ التسجيل'
        for entry in entries:
            row_cells = table.add_row().cells
            row_cells[0].text = entry.name
            row_cells[1].text = entry.identity
            row_cells[2].text = entry.payment_status
            row_cells[3].text = entry.whatsapp_name if entry.whatsapp_name else '--'
            row_cells[4].text = str(entry.entry_date)
    else:
        document.add_paragraph('لا توجد مدخلات جديدة لفك اليوم.')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="fak_report_{today.strftime("%Y-%m-%d")}.docx"'
    document.save(response)
    return response

