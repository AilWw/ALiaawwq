from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from .models import SocialSecurityEntry
from .forms import SocialSecurityEntryForm
from django.db.models import Q
import openpyxl
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
import datetime

def social_security_list(request):
    query = request.GET.get("q")
    entries = SocialSecurityEntry.objects.all()
    if query:
        entries = entries.filter(Q(name__icontains=query) | Q(identity__icontains=query))
    return render(request, 'social_security/social_security_list.html', {'entries': entries, 'query': query})

def social_security_create(request):
    if request.method == 'POST':
        form = SocialSecurityEntryForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('social_security_list')
    else:
        form = SocialSecurityEntryForm()
    return render(request, 'social_security/social_security_form.html', {'form': form})

def social_security_update(request, pk):
    entry = get_object_or_404(SocialSecurityEntry, pk=pk)
    if request.method == 'POST':
        form = SocialSecurityEntryForm(request.POST, instance=entry)
        if form.is_valid():
            form.save()
            return redirect('social_security_list')
    else:
        form = SocialSecurityEntryForm(instance=entry)
    return render(request, 'social_security/social_security_form.html', {'form': form})

def social_security_delete(request, pk):
    entry = get_object_or_404(SocialSecurityEntry, pk=pk)
    if request.method == 'POST':
        entry.delete()
        return redirect('social_security_list')
    return render(request, 'social_security/social_security_confirm_delete.html', {'entry': entry})

def export_social_security_to_excel(request):
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="social_security_data.xlsx"'

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Social Security Data"

    # Write headers
    columns = ['Name', 'Identity', 'Registration Date']
    sheet.append(columns)

    # Write data
    entries = SocialSecurityEntry.objects.all()
    for entry in entries:
        sheet.append([entry.name, entry.identity, entry.registration_date])

    workbook.save(response)
    return response

def generate_social_security_report(request):
    document = Document()
    document.add_heading('تقرير الضمان الاجتماعي اليومي', 0)

    today = datetime.date.today()
    document.add_paragraph(f'التاريخ: {today.strftime("%Y-%m-%d")}')

    entries = SocialSecurityEntry.objects.filter(registration_date=today)
    if entries:
        document.add_heading('مدخلات اليوم', level=1)
        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'الاسم'
        hdr_cells[1].text = 'الهوية'
        hdr_cells[2].text = 'تاريخ التسجيل'
        for entry in entries:
            row_cells = table.add_row().cells
            row_cells[0].text = entry.name
            row_cells[1].text = entry.identity
            row_cells[2].text = str(entry.registration_date)
    else:
        document.add_paragraph('لا توجد مدخلات جديدة للضمان الاجتماعي اليوم.')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="social_security_report_{today.strftime("%Y-%m-%d")}.docx"'
    document.save(response)
    return response

