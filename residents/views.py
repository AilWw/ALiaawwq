from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from .models import ResidentEntry
from .forms import ResidentEntryForm
from django.db.models import Q
import openpyxl
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
import datetime

def residents_list(request):
    query = request.GET.get("q")
    entries = ResidentEntry.objects.all()
    if query:
        entries = entries.filter(Q(name__icontains=query) | Q(residency_number__icontains=query) | Q(company__icontains=query))
    return render(request, 'residents/residents_list.html', {'entries': entries, 'query': query})

def residents_create(request):
    if request.method == 'POST':
        form = ResidentEntryForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('residents_list')
    else:
        form = ResidentEntryForm()
    return render(request, 'residents/residents_form.html', {'form': form})

def residents_update(request, pk):
    entry = get_object_or_404(ResidentEntry, pk=pk)
    if request.method == 'POST':
        form = ResidentEntryForm(request.POST, instance=entry)
        if form.is_valid():
            form.save()
            return redirect('residents_list')
    else:
        form = ResidentEntryForm(instance=entry)
    return render(request, 'residents/residents_form.html', {'form': form})

def residents_delete(request, pk):
    entry = get_object_or_404(ResidentEntry, pk=pk)
    if request.method == 'POST':
        entry.delete()
        return redirect('residents_list')
    return render(request, 'residents/residents_confirm_delete.html', {'entry': entry})

def export_residents_to_excel(request):
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="residents_data.xlsx"'

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Residents Data"

    # Write headers
    columns = ['Name', 'Residency Number', 'Company', 'Expiry Date']
    sheet.append(columns)

    # Write data
    entries = ResidentEntry.objects.all()
    for entry in entries:
        sheet.append([entry.name, entry.residency_number, entry.company, entry.expiry_date])

    workbook.save(response)
    return response

def generate_residents_report(request):
    document = Document()
    document.add_heading('تقرير المقيمين اليومي', 0)

    today = datetime.date.today()
    document.add_paragraph(f'التاريخ: {today.strftime("%Y-%m-%d")}')

    entries = ResidentEntry.objects.filter(expiry_date__gte=today) # Example: report on residents with upcoming expiry
    if entries:
        document.add_heading('مدخلات اليوم', level=1)
        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'الاسم'
        hdr_cells[1].text = 'رقم الإقامة'
        hdr_cells[2].text = 'الشركة'
        hdr_cells[3].text = 'تاريخ الانتهاء'
        for entry in entries:
            row_cells = table.add_row().cells
            row_cells[0].text = entry.name
            row_cells[1].text = entry.residency_number
            row_cells[2].text = entry.company
            row_cells[3].text = str(entry.expiry_date)
    else:
        document.add_paragraph('لا توجد مدخلات جديدة للمقيمين اليوم.')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="residents_report_{today.strftime("%Y-%m-%d")}.docx"'
    document.save(response)
    return response

